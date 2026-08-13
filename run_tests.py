"""
run_tests.py — Validates the PII redaction output
Run: python run_tests.py
"""

import os
import zipfile
from docx import Document

_HERE = os.path.dirname(os.path.abspath(__file__))
_REDACTED_PATH = os.path.join(_HERE, "Red_Herring_Prospectus_REDACTED.docx")

redacted = Document(_REDACTED_PATH)

# Collect ALL text from redacted doc (paragraphs + tables)
all_text = ' '.join(p.text for p in redacted.paragraphs)
all_text += ' '.join(
    cell.text for table in redacted.tables
    for row in table.rows for cell in row.cells
)

# ── TEST 1: PII must be GONE ─────────────────────────────────────────────────
known_pii = [
    'Sarthak.malvadkar@kshinterantional.com',
    'Sarthak Malvadkar',
    'ksh.ipo@nuvama.com',
    'customerservice.mb@nuvama.com',
    'Lokesh Shah',
    'Soumavo Sarkar',
    'Kishan Rastogi',
    'Cherag Gyara',
    'parag.pansare@kirtanepandit.com',
    'sheetal.parab@nuvama.com',
    'Sharmila Joshi',
    'Manisha Shukla',
    'Anand Soni',
    'Tushar Wakhele',
    'Ashish Mathew Pulloor',
    'Varun Badai',
    'Prakash Boricha',
    'Shanti Gopalkrishnan',
]

print('=' * 62)
print('  TEST 1 — PII MUST BE GONE FROM OUTPUT')
print('=' * 62)
p1 = f1 = 0
for pii in known_pii:
    if pii.lower() in all_text.lower():
        print(f'  FAIL ❌  {pii}')
        f1 += 1
    else:
        print(f'  PASS ✅  {pii}')
        p1 += 1
print(f'\n  Score: {p1}/{len(known_pii)} passed')

# ── TEST 2: Non-PII must be PRESERVED ────────────────────────────────────────
non_pii = [
    'Red Herring Prospectus',
    'SEBI ICDR Regulations',
    'Book Running Lead Managers',
    'KSH International Limited',
    'Registrar of Companies',
    'Maharashtra',
    'Equity Shares',
]

print()
print('=' * 62)
print('  TEST 2 — NON-PII MUST BE PRESERVED')
print('=' * 62)
p2 = f2 = 0
for phrase in non_pii:
    if phrase.lower() in all_text.lower():
        print(f'  PASS ✅  "{phrase}" — correctly kept')
        p2 += 1
    else:
        print(f'  WARN ⚠️   "{phrase}" — unexpectedly removed!')
        f2 += 1
print(f'\n  Score: {p2}/{len(non_pii)} preserved')

# ── TEST 3: Output file is a valid DOCX ──────────────────────────────────────
print()
print('=' * 62)
print('  TEST 3 — OUTPUT FILE VALIDITY')
print('=' * 62)
path = _REDACTED_PATH
size_kb = os.path.getsize(path) / 1024
print(f'  File size   : {size_kb:.1f} KB')
with zipfile.ZipFile(path) as z:
    has_doc = 'word/document.xml' in z.namelist()
print(f'  Valid DOCX  : {"YES - OK" if has_doc else "NO - BROKEN"}')
print(f'  Paragraphs  : {len(redacted.paragraphs)}')
print(f'  Tables      : {len(redacted.tables)}')

print()
print('=' * 62)
total = p1 + p2
total_possible = len(known_pii) + len(non_pii)
print(f'  OVERALL: {total}/{total_possible} checks passed')
print('=' * 62)
