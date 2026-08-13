"""
comprehensive_test.py
=====================
Extracts ALL real PII from the original Red Herring Prospectus,
then verifies every single instance is gone from the redacted output.

This is a data-driven test — no hand-picking. It finds PII programmatically.
"""

import re
import os
from docx import Document

BASE     = os.path.dirname(os.path.abspath(__file__))
ORIG     = os.path.join(BASE, "RHP_copy.docx")
REDACTED = os.path.join(BASE, "Red_Herring_Prospectus_REDACTED.docx")

# ─── Load both documents ──────────────────────────────────────────────────────
print("Loading documents...")
orig_doc     = Document(ORIG)
redacted_doc = Document(REDACTED)

def get_all_text(doc):
    """Extract all text from paragraphs + tables."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)

orig_text     = get_all_text(orig_doc)
redacted_text = get_all_text(redacted_doc)

# ─── PII Extractors for the ORIGINAL doc ─────────────────────────────────────

EMAIL_RE   = re.compile(r"(?<![/\w])[\w.+-]{1,64}@[\w.-]+\.[a-zA-Z]{2,6}(?!\w)", re.I)
PHONE_RE   = re.compile(r"(?:\+91[\s\-]?)(?:[\s\-]?\d){10}|0\d{2,4}[\s\-]\d{6,8}", re.I)
CONTACT_RE = re.compile(
    r"(?:Contact\s+Person|Contact\s+person|Compliance\s+Officer|Company\s+Secretary)"
    r"[:\s]+([A-Z][a-zA-Z]+(?: [A-Z][a-zA-Z]+){0,4}(?:\s*/\s*[A-Z][a-zA-Z]+(?: [A-Z][a-zA-Z]+){0,4})*)"
)
NAME_TOKEN = re.compile(r"[A-Z][a-z]+(?: [A-Z][a-z]+){0,4}")

# Words that look like Title Case names but are NOT person names
_NAME_STOPWORDS = {
    "email", "website", "telephone", "contact", "person", "registration",
    "no", "number", "sebi", "india", "limited", "private", "public",
    "floor", "tower", "block", "wing", "building", "road", "street",
    "marg", "nagar", "east", "west", "north", "south", "mumbai", "pune",
    "delhi", "maharashtra", "bandra", "kurla", "complex",
}

def _is_valid_name(name: str) -> bool:
    """Returns True only if the name looks like an actual person name."""
    words = name.lower().split()
    # Reject if any word is a known non-name
    if any(w in _NAME_STOPWORDS for w in words):
        return False
    # Reject single-word matches shorter than 4 chars
    if len(words) == 1 and len(name) < 4:
        return False
    return True

# Extract all emails from original
orig_emails = sorted(set(m.group().strip() for m in EMAIL_RE.finditer(orig_text)))

# Extract all phone numbers from original
orig_phones = sorted(set(m.group().strip() for m in PHONE_RE.finditer(orig_text)))

# Extract all contact-person names from original
orig_names = set()
for m in CONTACT_RE.finditer(orig_text):
    for name_m in NAME_TOKEN.finditer(m.group(1)):
        name = name_m.group().strip()
        if len(name) > 3 and _is_valid_name(name):
            orig_names.add(name)
orig_names = sorted(orig_names)

# ─── Report what was found ────────────────────────────────────────────────────
print(f"\nFound in ORIGINAL document:")
print(f"  Emails      : {len(orig_emails)}")
print(f"  Phone numbers: {len(orig_phones)}")
print(f"  Person names : {len(orig_names)}")

# ─── TEST: None of the original PII should exist in redacted ─────────────────

def run_test(label, pii_list, category):
    passed = failed = 0
    failures = []
    for pii in pii_list:
        if pii.lower() in redacted_text.lower():
            failures.append(pii)
            failed += 1
        else:
            passed += 1
    return passed, failed, failures

print("\n" + "=" * 66)
print("  COMPREHENSIVE PII REDACTION TEST")
print("  (All real PII extracted from original document)")
print("=" * 66)

# ── Emails ────────────────────────────────────────────────────────────────────
ep, ef, email_fails = run_test("EMAIL", orig_emails, "emails")
print(f"\n📧 EMAILS  —  {ep}/{len(orig_emails)} redacted")
if email_fails:
    print("  FAILED (still visible):")
    for f in email_fails:
        print(f"    ❌ {f}")
else:
    print("  All emails successfully redacted ✅")

# ── Phones ────────────────────────────────────────────────────────────────────
pp, pf, phone_fails = run_test("PHONE", orig_phones, "phones")
print(f"\n📱 PHONES  —  {pp}/{len(orig_phones)} redacted")
if phone_fails:
    print("  FAILED (still visible):")
    for f in phone_fails:
        print(f"    ❌ {f}")
else:
    print("  All phone numbers successfully redacted ✅")

# ── Person Names ──────────────────────────────────────────────────────────────
np_, nf, name_fails = run_test("PERSON", orig_names, "names")
print(f"\n👤 NAMES   —  {np_}/{len(orig_names)} redacted")
if name_fails:
    print("  FAILED (still visible):")
    for f in name_fails:
        print(f"    ❌ {f}")
else:
    print("  All contact-person names successfully redacted ✅")

# ── Non-PII Preservation Check ────────────────────────────────────────────────
preserve_phrases = [
    "Red Herring Prospectus",
    "SEBI ICDR Regulations",
    "KSH International Limited",
    "Book Running Lead Managers",
    "Registrar of Companies",
    "Maharashtra",
    "Equity Shares",
    "Fresh Issue",
    "Offer for Sale",
    "Board of Directors",
]
pp2 = pf2 = 0
preserve_fails = []
for phrase in preserve_phrases:
    if phrase.lower() in redacted_text.lower():
        pp2 += 1
    else:
        pf2 += 1
        preserve_fails.append(phrase)

print(f"\n🏛️  NON-PII PRESERVED  —  {pp2}/{len(preserve_phrases)} intact")
if preserve_fails:
    for f in preserve_fails:
        print(f"  OVER-REDACTED ⚠️  {f}")
else:
    print("  All non-PII text correctly preserved ✅")

# ── Final Scorecard ───────────────────────────────────────────────────────────
total_pii     = len(orig_emails) + len(orig_phones) + len(orig_names)
total_passed  = ep + pp + np_
total_failed  = ef + pf + nf
total_non_pii = len(preserve_phrases)

print()
print("=" * 66)
print("  FINAL SCORECARD")
print("=" * 66)
print(f"  PII Redacted      : {total_passed}/{total_pii}  ({100*total_passed/total_pii:.1f}%)")
print(f"  PII Missed        : {total_failed}")
print(f"  Non-PII Preserved : {pp2}/{total_non_pii}")
print(f"  Over-Redacted     : {pf2}")
recall    = total_passed / total_pii if total_pii > 0 else 0
precision = total_passed / (total_passed + pf2) if (total_passed + pf2) > 0 else 0
f1        = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0
print()
print(f"  Recall      : {recall:.2%}")
print(f"  Precision   : {precision:.2%}")
print(f"  F1 Score    : {f1:.2%}")
print("=" * 66)

# ── Full PII inventory ────────────────────────────────────────────────────────
print("\n\nFULL PII INVENTORY FROM ORIGINAL DOCUMENT:")
print("-" * 50)
print(f"\nAll {len(orig_emails)} emails found:")
for e in orig_emails:
    status = "❌ STILL IN REDACTED" if e.lower() in redacted_text.lower() else "✅"
    print(f"  {status}  {e}")

print(f"\nAll {len(orig_phones)} phone numbers found:")
for p in orig_phones:
    status = "❌ STILL IN REDACTED" if p.lower() in redacted_text.lower() else "✅"
    print(f"  {status}  {p}")

print(f"\nAll {len(orig_names)} contact-person names found:")
for n in orig_names:
    status = "❌ STILL IN REDACTED" if n.lower() in redacted_text.lower() else "✅"
    print(f"  {status}  {n}")
