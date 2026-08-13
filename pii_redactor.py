"""
PII Redaction Tool
------------------
Reads a .docx file, finds all personally identifiable information,
replaces each unique value with a consistent fake alternative, and
saves the result as a new .docx.

Approach:
  - Layer 1: regex for structured PII (emails, phones, PAN, IP, SSN, DOB)
  - Layer 2: spaCy en_core_web_lg for names and org mentions
  - Fallback: keyword-context regex for names that NER misses
    (e.g., "Contact Person: Lokesh Shah / Soumavo Sarkar")

Each unique real value always maps to the same fake value throughout
the document — so the output stays internally consistent.

Known tradeoffs:
  - spaCy sometimes misses names embedded in dense legal prose
  - ORG entity is broad; we whitelist common financial/regulatory terms
    to avoid over-redacting things like "Mutual Funds" or "Registrar"
  - Financial dates (Q1 FY2025 etc.) are intentionally NOT touched;
    only birth-context dates get redacted

Dependencies:
    pip install python-docx spacy faker
    python -m spacy download en_core_web_lg
"""

import os
import re
import random
import logging
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

import spacy
from docx import Document
from faker import Faker


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)

# ── paths ──────────────────────────────────────────────────────────────────
_here = os.path.dirname(os.path.abspath(__file__))
# prefer the unlocked copy if Word has the original open
_src  = os.path.join(_here, "RHP_copy.docx")
if not os.path.exists(_src):
    _src = os.path.join(_here, "Red Herring Prospectus.docx")

INPUT_PATH  = _src
OUTPUT_PATH = os.path.join(_here, "Red_Herring_Prospectus_REDACTED.docx")

# fixed seed so every run produces the same fake values (reproducible)
_SEED = 42
random.seed(_SEED)
Faker.seed(_SEED)
fake = Faker("en_IN")

nlp: Optional[spacy.language.Language] = None   # lazy-loaded


# ── replacement registry ───────────────────────────────────────────────────
# Maps real PII → fake value. Populated on first encounter, reused after.
# This is what keeps "Lokesh Shah" → "Rajesh Verma" everywhere in the doc.

@dataclass
class _Registry:
    names:    Dict[str, str] = field(default_factory=dict)
    emails:   Dict[str, str] = field(default_factory=dict)
    phones:   Dict[str, str] = field(default_factory=dict)
    orgs:     Dict[str, str] = field(default_factory=dict)
    misc:     Dict[str, str] = field(default_factory=dict)   # IPs, PANs, SSNs …
    counts:   Dict[str, int] = field(default_factory=lambda: {
        "names": 0, "emails": 0, "phones": 0, "orgs": 0, "misc": 0
    })

_reg = _Registry()


def _memo(store: dict, key: str, gen) -> str:
    """Return cached fake value, or generate + cache a new one."""
    k = key.strip()
    if k not in store:
        store[k] = gen(k)
    return store[k]


# ── fake value generators ──────────────────────────────────────────────────

def _fake_name(_):
    return fake.name()

def _fake_email(_):
    return f"{fake.user_name()}@{fake.domain_name()}"

def _fake_phone(_):
    digits = "".join(str(random.randint(0, 9)) for _ in range(10))
    return f"+91 {digits}"

def _fake_org(_):
    return fake.company()

def _fake_ip(_):
    return fake.ipv4_private()

def _fake_pan(_):
    alpha = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    return (
        "".join(random.choice(alpha) for _ in range(5))
        + "".join(str(random.randint(0, 9)) for _ in range(4))
        + random.choice(alpha)
    )

def _fake_dob(_):
    return fake.date_of_birth(minimum_age=25, maximum_age=75).strftime("%d/%m/%Y")


# ── regex rules ────────────────────────────────────────────────────────────
# Tuples of (label, pattern, replacement_fn).
# Order matters — email must come before the generic phone sweep.

_REGEX_RULES: List[Tuple[str, re.Pattern, callable]] = [

    ("email",
     re.compile(r"(?<![/\w])[\w.+-]{1,64}@[\w.-]+\.[a-zA-Z]{2,6}(?!\w)", re.I),
     lambda m: _memo(_reg.emails, m.group(), _fake_email)),

    # +91 style mobile
    ("phone",
     re.compile(r"(?<!\d)(?:\+\s*91[\s\-]?)(?:[\s\-]?\d){10}(?!\d)", re.I),
     lambda m: _memo(_reg.phones, re.sub(r"\D", "", m.group()), _fake_phone)),

    # STD landline  e.g. 022-40094400
    ("phone_land",
     re.compile(r"\b0\d{2,4}[\s\-]\d{6,8}\b"),
     lambda m: _memo(_reg.phones, re.sub(r"\D", "", m.group()), _fake_phone)),

    # PAN card  e.g. ABCDE1234F
    ("pan",
     re.compile(r"\b[A-Z]{5}[0-9]{4}[A-Z]\b"),
     lambda m: _memo(_reg.misc, m.group(), _fake_pan)),

    # US-style SSN  (just in case)
    ("ssn",
     re.compile(r"\b\d{3}-\d{2}-\d{4}\b"),
     lambda m: _memo(_reg.misc, m.group(), lambda _: fake.ssn())),

    # IPv4
    ("ip",
     re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b"),
     lambda m: _memo(_reg.misc, m.group(), _fake_ip)),

    # Explicit date-of-birth patterns only — not generic financial dates
    ("dob",
     re.compile(
         r"(?:date\s+of\s+birth|dob|born\s+on)\s*[:\-]?\s*\d{1,2}[/\-.]\d{1,2}[/\-.]\d{2,4}",
         re.I,
     ),
     lambda m: "Date of Birth: " + _memo(_reg.misc, m.group(), _fake_dob)),
]


# ── whitelist — things NER might flag but should never be redacted ──────────
# spaCy's ORG entity is overly broad in legal documents.
_WHITELIST = {
    "SEBI", "RBI", "BSE", "NSE", "NSDL", "CDSL", "ROC", "GOI",
    "INDIA", "MAHARASHTRA", "MUMBAI", "PUNE", "DELHI",
    "UPI", "ASBA", "QIB", "IPO", "PAN", "KYC",
    "REGISTRAR OF COMPANIES", "REGISTRAR",
    "SUPREME COURT", "HIGH COURT", "INCOME TAX",
    "MUTUAL FUNDS", "LIFE INSURANCE", "PENSION FUNDS",
    "EQUITY SHARES", "PREFERENCE SHARES",
    "BOOK RUNNING LEAD MANAGERS", "LEAD MANAGERS",
    "ANCHOR INVESTORS", "RETAIL INDIVIDUAL INVESTORS",
    "NON-INSTITUTIONAL INVESTORS",
}

def _whitelisted(text: str) -> bool:
    return text.strip().upper() in _WHITELIST or len(text.strip()) <= 2


# ── keyword-context name extractor (NER fallback) ──────────────────────────
# Catches person names that appear right after a label like "Contact Person:".
# Handles slash-separated lists: "Lokesh Shah / Soumavo Sarkar"

_CP_PATTERN = re.compile(
    r"(?:Contact\s+Person|Compliance\s+Officer|Company\s+Secretary)"
    r"[:\s]+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,4}"
    r"(?:\s*/\s*[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,4})*)",
)
_NAME_TOKEN = re.compile(r"[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3}")


def _keyword_names(text: str) -> List[Tuple[int, int, str]]:
    """Find person names via keyword context — returns (start, end, fake) tuples."""
    hits = []
    for m in _CP_PATTERN.finditer(text):
        base = m.start(1)
        for nm in _NAME_TOKEN.finditer(m.group(1)):
            name = nm.group().strip()
            if _whitelisted(name) or len(name) < 4:
                continue
            fake_val = _memo(_reg.names, name, _fake_name)
            _reg.counts["names"] += 1
            hits.append((base + nm.start(), base + nm.end(), fake_val))
    return hits


# ── core redaction functions ────────────────────────────────────────────────

def _apply_regex(text: str) -> str:
    for label, pattern, replacer in _REGEX_RULES:
        def _sub(m, _lbl=label, _rep=replacer):
            val = _rep(m)
            cat = "emails" if _lbl == "email" else "phones" if "phone" in _lbl else "misc"
            _reg.counts[cat] += 1
            return val
        text = pattern.sub(_sub, text)
    return text


def _apply_ner(text: str) -> str:
    """Run spaCy NER + keyword fallback on the text."""
    global nlp
    if nlp is None:
        log.info("Loading spaCy model …")
        nlp = spacy.load("en_core_web_lg")

    doc = nlp(text)
    replacements: List[Tuple[int, int, str]] = []
    covered: set = set()

    for ent in doc.ents:
        if _whitelisted(ent.text):
            continue

        if ent.label_ == "PERSON":
            fake_val = _memo(_reg.names, ent.text.strip(), _fake_name)
            _reg.counts["names"] += 1
            replacements.append((ent.start_char, ent.end_char, fake_val))
            covered.add((ent.start_char, ent.end_char))

        elif ent.label_ == "ORG":
            orig = ent.text.strip()
            # skip all-caps abbreviations — those are almost never personal ORGs
            if len(orig) > 3 and not orig.isupper():
                fake_val = _memo(_reg.orgs, orig, _fake_org)
                _reg.counts["orgs"] += 1
                replacements.append((ent.start_char, ent.end_char, fake_val))
                covered.add((ent.start_char, ent.end_char))

    # keyword fallback for names NER missed
    for start, end, fake_val in _keyword_names(text):
        overlap = any(not (end <= s or start >= e) for s, e in covered)
        if not overlap:
            replacements.append((start, end, fake_val))

    # apply in reverse so char offsets stay valid
    chars = list(text)
    for start, end, fv in sorted(replacements, key=lambda x: x[0], reverse=True):
        chars[start:end] = list(fv)

    return "".join(chars)


def redact(text: str) -> str:
    """
    Main pipeline:
      1. NER on original text first — preserves context for name detection
      2. Regex on the NER-cleaned text — handles structured PII
    Running NER before regex means a phone number replacement won't confuse
    spaCy's context window for the person name next to it.
    """
    if not text or not text.strip():
        return text
    return _apply_regex(_apply_ner(text))


# ── document processing ────────────────────────────────────────────────────

def _redact_para(para) -> None:
    """
    Redact a single paragraph in place.
    DOCX splits paragraphs into 'runs' (each with its own bold/italic/font
    style). We merge all run text, redact it as one string, then put the
    result into run[0] and clear the rest — preserving formatting.
    """
    if not para.runs:
        return

    original = "".join(r.text for r in para.runs)
    if not original.strip():
        return

    cleaned = redact(original)
    if cleaned == original:
        return

    para.runs[0].text = cleaned
    for r in para.runs[1:]:
        r.text = ""


def process(input_path: str, output_path: str) -> None:
    log.info("Opening: %s", input_path)
    doc = Document(input_path)

    log.info("Scanning %d paragraphs …", len(doc.paragraphs))
    for para in doc.paragraphs:
        _redact_para(para)

    log.info("Scanning %d tables …", len(doc.tables))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _redact_para(para)

    # headers and footers often contain company name / contact info
    for section in doc.sections:
        for hf in [
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ]:
            if hf is None:
                continue
            for para in hf.paragraphs:
                _redact_para(para)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _redact_para(para)

    doc.save(output_path)
    log.info("Saved → %s", output_path)


def _print_summary() -> None:
    total = sum(_reg.counts.values())
    log.info("─" * 45)
    log.info("Redaction summary")
    log.info("─" * 45)
    for cat, n in _reg.counts.items():
        log.info("  %-8s  %d", cat, n)
    log.info("  %-8s  %d", "total", total)
    log.info("─" * 45)


if __name__ == "__main__":
    process(INPUT_PATH, OUTPUT_PATH)
    _print_summary()
    print(f"\nDone. Output → {OUTPUT_PATH}")
