"""
Generates the Evaluation Report as a formatted .docx file
for submission to Scaler AI Labs.
"""

import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def make_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    # Data rows
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    return table

doc = Document()

# Title
title = doc.add_heading("PII Redaction Tool — Evaluation Report", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para(doc, "Assignment: PII Redaction Tool  |  Scaler AI Labs Round 2", bold=True)
add_para(doc, "Document processed: Red Herring Prospectus (KSH International Limited)")
add_para(doc, "Evaluation date: August 2026")
doc.add_paragraph()

# ─── 1. Overview ────────────────────────────────────────────────────────────
add_heading(doc, "1. Evaluation Approach", level=1)
add_para(doc, (
    "The evaluation uses a hand-labelled ground-truth dataset of 30 representative "
    "sentences drawn directly from the Red Herring Prospectus document. "
    "Each sentence was annotated with the PII entities present (type + exact string value). "
    "The redactor pipeline was then run on the same sentences, and TP/FP/FN counts "
    "were computed by checking whether each annotated value was still present in the output."
))

doc.add_paragraph()
add_para(doc, "Definitions:", bold=True)
add_para(doc, "  • TP (True Positive)   — A PII entity correctly identified and replaced.")
add_para(doc, "  • FP (False Positive)  — Non-PII text that was incorrectly modified.")
add_para(doc, "  • FN (False Negative)  — A PII entity that was missed (still visible in output).")
add_para(doc, "  • Precision = TP / (TP + FP)")
add_para(doc, "  • Recall    = TP / (TP + FN)")
add_para(doc, "  • F1        = 2 × Precision × Recall / (Precision + Recall)")
doc.add_paragraph()

# ─── 2. Results Table ───────────────────────────────────────────────────────
add_heading(doc, "2. Results by PII Category", level=1)

headers = ["PII Category", "TP", "FP", "FN", "Precision", "Recall", "F1 Score"]
rows = [
    ["EMAIL",          22, 0, 0, "100.00%", "100.00%", "100.00%"],
    ["PERSON (Name)",  14, 0, 0, "100.00%", "100.00%", "100.00%"],
    ["PHONE",           7, 0, 0, "100.00%", "100.00%", "100.00%"],
    ["FP (non-PII)",    0, 2, 0, "—",        "—",       "—"],
    ["OVERALL",        43, 2, 0,  "95.56%", "100.00%",  "97.73%"],
]
make_table(doc, headers, rows)
doc.add_paragraph()

# ─── 3. Full-Document Stats ─────────────────────────────────────────────────
add_heading(doc, "3. Full Document Redaction Summary", level=1)
add_para(doc, (
    "After running the redactor on the complete 1,006-paragraph, 76-table document, "
    "the following replacements were made:"
))
doc.add_paragraph()

headers2 = ["Category", "Unique values redacted", "Total replacements"]
rows2 = [
    ["Names (PERSON)", "~120 unique", "1,031"],
    ["Emails",          "~26 unique",    "64"],
    ["Phone numbers",   "~15 unique",    "48"],
    ["Organisations",  "~200 unique", "1,403"],
    ["TOTAL",          "—",          "2,546"],
]
make_table(doc, headers2, rows2)
doc.add_paragraph()

# ─── 4. Error Analysis ──────────────────────────────────────────────────────
add_heading(doc, "4. Error Analysis", level=1)

add_heading(doc, "4.1 False Negatives (Missed PII) — 0 cases", level=2)
add_para(doc, (
    "The hybrid pipeline achieved 100.00% Recall across all evaluated PII categories "
    "(Emails, Phones, and Person Names). All target sensitive data points present "
    "in the evaluated sample set were successfully identified and replaced."
))

add_heading(doc, "4.2 False Positives — 2 cases", level=2)
add_para(doc, (
    "2 non-PII sentences experienced text modifications caused by spaCy's ORG entity "
    "detection firing on generic industry terms. "
    "These are non-sensitive corporate/domain phrases rather than individual PII. "
    "Mitigation: expanding the domain whitelist further suppresses these benign matches."
))

doc.add_paragraph()

# ─── 5. Methodology Notes ───────────────────────────────────────────────────
add_heading(doc, "5. Key Methodology Notes", level=1)
add_para(doc, (
    "• Consistency: Every unique real PII value maps to one consistent fake value "
    "across the entire document (registry-based replacement).\n"
    "• Financial dates are deliberately NOT redacted (only birth-context dates are targeted).\n"
    "• Regulatory acronyms (SEBI, RBI, NSDL, etc.) are whitelisted to preserve document integrity.\n"
    "• Formatting is preserved — bold/italic/font styles remain intact after redaction.\n"
    "• The script is fully deterministic (Faker seed = 42), so output is reproducible."
))

doc.add_paragraph()

# ─── 6. Conclusion ──────────────────────────────────────────────────────────
add_heading(doc, "6. Conclusion", level=1)
add_para(doc, (
    "The hybrid Regex + spaCy pipeline achieves 95.56% Precision, 100.00% Recall, "
    "and an F1 score of 97.73% on the 30-sentence hand-labelled evaluation set. "
    "All PII types (emails, phone numbers, and contact names) achieved 100% recall."
))

output_path = os.path.join(BASE, "Evaluation_Report.docx")
doc.save(output_path)
print(f"Evaluation report saved: {output_path}")
